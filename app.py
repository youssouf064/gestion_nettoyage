import io
import csv
import os
import re
from datetime import datetime
import openpyxl
import pdfplumber
from pypdf import PdfReader
import docx
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, jsonify, session, Response

app = Flask(__name__)
app.secret_key = "cle_secrete_super_securisee_youssouf"

# --- CONFIGURATION HYBRIDE (NEON SUR RENDER / MYSQL EN LOCAL) ---
IS_RENDER = 'RENDER' in os.environ

if IS_RENDER:
    import psycopg2
    DATABASE_URL = os.environ.get('DATABASE_URL')
    if DATABASE_URL and DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)
else:
    import mysql.connector
    DB_CONFIG_LOCAL = {
        'host': 'localhost',
        'user': 'nettoyage_user',
        'password': 'MonMotDePasseSecurise123!',
        'database': 'gestion_nettoyage'
    }

def get_db_connection():
    if IS_RENDER:
        return psycopg2.connect(DATABASE_URL)
    else:
        return mysql.connector.connect(**DB_CONFIG_LOCAL)

def init_db():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if IS_RENDER:
        cursor.execute('''CREATE TABLE IF NOT EXISTS sites (
            id SERIAL PRIMARY KEY, 
            nom VARCHAR(255) NOT NULL, 
            adresse VARCHAR(255),
            latitude REAL,
            longitude REAL
        )''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS employes (
            matricule VARCHAR(100) PRIMARY KEY, 
            nom VARCHAR(255) NOT NULL, 
            prenom VARCHAR(255) NOT NULL, 
            salaire_base REAL NOT NULL, 
            statut VARCHAR(50) DEFAULT 'Actif',
            id_site_affecte INT,
            FOREIGN KEY(id_site_affecte) REFERENCES sites(id) ON DELETE SET NULL
        )''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS pointages (
            id SERIAL PRIMARY KEY,
            matricule_employe VARCHAR(100),
            id_site INT,
            date_jour VARCHAR(50) NOT NULL,
            heure_arrivee VARCHAR(50),
            heure_depart VARCHAR(50),
            FOREIGN KEY(matricule_employe) REFERENCES employes(matricule) ON DELETE CASCADE,
            FOREIGN KEY(id_site) REFERENCES sites(id) ON DELETE CASCADE
        )''')
        conn.commit()
        
        # Mises à jour des colonnes facultatives avec rollback sécurisé
        colonnes_ajouts = [
            "ALTER TABLE pointages ADD COLUMN IF NOT EXISTS latitude REAL;",
            "ALTER TABLE pointages ADD COLUMN IF NOT EXISTS longitude REAL;",
            "ALTER TABLE employes ADD COLUMN IF NOT EXISTS equipe VARCHAR(50) DEFAULT 'MATIN';",
            "ALTER TABLE pointages ADD COLUMN IF NOT EXISTS equipe VARCHAR(50) DEFAULT 'MATIN';",
            "ALTER TABLE pointages ADD COLUMN IF NOT EXISTS fiche_id VARCHAR(100);",
            "ALTER TABLE pointages ADD COLUMN IF NOT EXISTS agent VARCHAR(255);",
            "ALTER TABLE pointages ADD COLUMN IF NOT EXISTS observations TEXT;"
        ]
        for req in colonnes_ajouts:
            try:
                cursor.execute(req)
                conn.commit()
            except Exception as e:
                conn.rollback()
                print(f"Note SQL (PostgreSQL) : {e}")

    else:
        cursor.execute('''CREATE TABLE IF NOT EXISTS sites (
            id INT AUTO_INCREMENT PRIMARY KEY, 
            nom VARCHAR(255) NOT NULL, 
            adresse VARCHAR(255),
            latitude REAL,
            longitude REAL
        )''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS employes (
            matricule VARCHAR(100) PRIMARY KEY, 
            nom VARCHAR(255) NOT NULL, 
            prenom VARCHAR(255) NOT NULL, 
            salaire_base REAL NOT NULL, 
            statut VARCHAR(50) DEFAULT 'Actif',
            id_site_affecte INT,
            FOREIGN KEY(id_site_affecte) REFERENCES sites(id) ON DELETE SET NULL
        )''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS pointages (
            id INT AUTO_INCREMENT PRIMARY KEY,
            matricule_employe VARCHAR(100),
            id_site INT,
            date_jour VARCHAR(50) NOT NULL,
            heure_arrivee VARCHAR(50),
            heure_depart VARCHAR(50),
            FOREIGN KEY(matricule_employe) REFERENCES employes(matricule) ON DELETE CASCADE,
            FOREIGN KEY(id_site) REFERENCES sites(id) ON DELETE CASCADE
        )''')
        conn.commit()
        
        colonnes_mysql = [
            "ALTER TABLE pointages ADD COLUMN latitude REAL NULL;",
            "ALTER TABLE pointages ADD COLUMN longitude REAL NULL;",
            "ALTER TABLE employes ADD COLUMN equipe VARCHAR(50) DEFAULT 'MATIN';",
            "ALTER TABLE pointages ADD COLUMN equipe VARCHAR(50) DEFAULT 'MATIN';",
            "ALTER TABLE pointages ADD COLUMN fiche_id VARCHAR(100) NULL;",
            "ALTER TABLE pointages ADD COLUMN agent VARCHAR(255) NULL;",
            "ALTER TABLE pointages ADD COLUMN observations TEXT NULL;"
        ]
        for req in colonnes_mysql:
            try:
                cursor.execute(req)
            except Exception:
                pass
        conn.commit()

    cursor.close()
    conn.close()

init_db()

CODE_SECRET_ADMIN = "1234"

# --- ACCÈS ET CONTRÔLE DU BUREAU PRINCIPAL ---

@app.route('/', methods=['GET', 'POST'])
def dashboard():
    if request.method == 'POST':
        code_saisi = request.form.get('code_admin')
        if code_saisi == CODE_SECRET_ADMIN:
            session['est_admin'] = True  
        else:
            return render_template('connexion_admin.html', erreur="Code incorrect. Accès refusé.")

    if not session.get('est_admin'):
        return render_template('connexion_admin.html', erreur=None)

    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT COUNT(*) FROM employes WHERE statut = 'Actif'")
    total_actifs = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM employes WHERE statut = 'En congé'")
    total_conges = cursor.fetchone()[0]
    
    if IS_RENDER:
        fonction_date = "CURRENT_DATE::text"
    else:
        fonction_date = "CURDATE()"

    cursor.execute(f'''
        SELECT s.id, s.nom, s.adresse,
        (SELECT COUNT(*) FROM pointages p WHERE p.id_site = s.id AND p.date_jour = {fonction_date} AND p.heure_depart IS NULL) as presents,
        (SELECT COUNT(*) FROM employes e WHERE e.id_site_affecte = s.id) as total_employes
        FROM sites s
    ''')
    liste_sites = cursor.fetchall()
    
    cursor.execute('''
        SELECT e.matricule, e.nom, e.prenom, e.salaire_base, e.statut, s.nom, e.equipe 
        FROM employes e
        LEFT JOIN sites s ON e.id_site_affecte = s.id
    ''')
    liste_employes = cursor.fetchall()
    
    cursor.execute("SELECT matricule, nom, prenom FROM employes WHERE statut = 'En congé'")
    employes_en_conge = cursor.fetchall()
    
    cursor.execute('''
        SELECT p.id, e.prenom, e.nom, s.nom, p.date_jour, p.heure_arrivee, p.heure_depart, p.latitude, p.longitude, COALESCE(p.equipe, 'MATIN') as equipe_p, p.fiche_id, p.agent, p.observations
        FROM pointages p
        JOIN employes e ON p.matricule_employe = e.matricule
        JOIN sites s ON p.id_site = s.id
        ORDER BY p.id DESC
    ''')
    historique_pointages = cursor.fetchall()
    
    cursor.close()
    conn.close()
    
    return render_template('dashboard.html', 
                           total_actifs=total_actifs, 
                           total_conges=total_conges, 
                           sites=liste_sites,
                           employes=liste_employes,
                           conges=employes_en_conge,
                           pointages=historique_pointages)


@app.route('/deconnexion')
def deconnexion():
    session.pop('est_admin', None)
    return redirect(url_for('espace_pointage'))


@app.route('/ajouter_site', methods=['POST'])
def ajouter_site():
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    nom_site = request.form.get('nom_site').strip()
    adresse_site = request.form.get('adresse_site').strip()
    if nom_site:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("INSERT INTO sites (nom, adresse) VALUES (%s, %s)", (nom_site, adresse_site))
        conn.commit()
        cursor.close()
        conn.close()
    return redirect(url_for('dashboard'))


@app.route('/ajouter_employe', methods=['POST'])
def ajouter_employe():
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    matricule = request.form.get('matricule').upper().strip()
    nom = request.form.get('nom').strip()
    prenom = request.form.get('prenom').strip()
    salaire = request.form.get('salaire')
    statut = request.form.get('statut')
    site_id = request.form.get('site_id')
    equipe = request.form.get('equipe', 'MATIN').upper().strip()
    
    if matricule and nom and prenom and salaire:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO employes (matricule, nom, prenom, salaire_base, statut, id_site_affecte, equipe)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        ''', (matricule, nom, prenom, float(salaire), statut, int(site_id), equipe))
        conn.commit()
        cursor.close()
        conn.close()
    return redirect(url_for('dashboard'))


@app.route('/importer_employes_csv', methods=['POST'])
def importer_employes_csv():
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    file = request.files.get('file_excel')
    if not file or file.filename == '':
        return f"<h3>Aucun fichier sélectionné.</h3><br><a href='/'>Retour</a>"

    filename = file.filename.lower()
    lignes_donnees = []
    texte_complet_pdf = ""

    try:
        if filename.endswith(('.xlsx', '.xls', '.csv')):
            try:
                if filename.endswith('.csv'):
                    df = pd.read_csv(file, sep=None, engine='python', dtype=str)
                else:
                    df = pd.read_excel(file, dtype=str)
                
                df = df.fillna('')
                lignes_donnees = [df.columns.astype(str).tolist()] + df.values.tolist()
            except Exception:
                file.seek(0)
                content = file.stream.read().decode("utf-8-sig", errors='ignore')
                stream = io.StringIO(content)
                separateur = ';' if ';' in content.split('\n')[0] else ','
                reader = csv.reader(stream, delimiter=separateur)
                lignes_donnees = list(reader)

        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        lignes_donnees.append(cells)

        elif filename.endswith('.pdf'):
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    text_p = page.extract_text() or ""
                    texte_complet_pdf += text_p + "\n"
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if any(row):
                                lignes_donnees.append([str(cell or '').strip() for cell in row])
                    
                    if not lignes_donnees and text_p:
                        for line in text_p.split('\n'):
                            if line.strip():
                                lignes_donnees.append([line.strip()])

        if not lignes_donnees and not texte_complet_pdf:
            return f"<h3>Le fichier est vide ou n'a pas pu être lu.</h3><br><a href='/'>Retour</a>"

        # Extraction des métadonnées dynamique (fiche_id, agent, observations)
        fiche_id = None
        agent_nom = None
        observations = None

        m_fiche = re.search(r'(?:FICHE|N[°º]|REF)\s*[:#-]?\s*([A-Z0-9-]+)', texte_complet_pdf, re.IGNORECASE)
        if m_fiche:
            fiche_id = m_fiche.group(1).strip()

        m_agent = re.search(r'(?:AGENT|RESPONSABLE|SUPERVISEUR)\s*[:#-]?\s*([A-Za-zÀ-ÿ\s-]+)', texte_complet_pdf, re.IGNORECASE)
        if m_agent:
            agent_nom = m_agent.group(1).split('\n')[0].strip()

        m_obs = re.search(r'(?:REMARQUES|OBSERVATIONS|NOTE)\s*[:#-]?\s*(.*)', texte_complet_pdf, re.IGNORECASE | re.DOTALL)
        if m_obs:
            observations = m_obs.group(1).strip()

        headers = [str(h).lower().strip() for h in lignes_donnees[0]] if lignes_donnees else []

        idx_mat = next((i for i, h in enumerate(headers) if 'matricule' in h), -1)
        idx_nom = next((i for i, h in enumerate(headers) if 'nom' in h and 'prenom' not in h), -1)
        idx_prenom = next((i for i, h in enumerate(headers) if 'prenom' in h), -1)
        idx_salaire = next((i for i, h in enumerate(headers) if 'salaire' in h), -1)
        idx_statut = next((i for i, h in enumerate(headers) if 'statut' in h), -1)
        idx_site = next((i for i, h in enumerate(headers) if 'site' in h or 'affectation' in h), -1)
        idx_equipe = next((i for i, h in enumerate(headers) if 'equipe' in h or 'équipe' in h), -1)

        conn = get_db_connection()
        cursor = conn.cursor()
        nb_importes = 0

        if idx_nom != -1 or idx_prenom != -1:
            idx_mat = 0 if idx_mat == -1 else idx_mat
            idx_nom = 1 if idx_nom == -1 else idx_nom
            idx_prenom = 2 if idx_prenom == -1 else idx_prenom
            idx_salaire = 3 if idx_salaire == -1 else idx_salaire

            for row in lignes_donnees[1:]:
                if len(row) <= max(idx_mat, idx_nom, idx_prenom):
                    continue

                matricule = str(row[idx_mat]).upper().strip()
                nom = str(row[idx_nom]).strip()
                prenom = str(row[idx_prenom]).strip()

                salaire_str = str(row[idx_salaire]) if idx_salaire < len(row) else '0'
                statut = str(row[idx_statut]) if (idx_statut != -1 and idx_statut < len(row)) else 'Actif'
                nom_site = str(row[idx_site]) if (idx_site != -1 and idx_site < len(row)) else ''
                equipe = str(row[idx_equipe]).upper().strip() if (idx_equipe != -1 and idx_equipe < len(row)) else 'MATIN'

                if not matricule or not nom or not prenom:
                    continue

                cursor.execute("SELECT matricule FROM employes WHERE matricule = %s", (matricule,))
                if cursor.fetchone():
                    continue

                id_site = get_or_create_site(cursor, nom_site)

                try:
                    salaire_flt = float(salaire_str.replace(',', '.').replace(' ', ''))
                except ValueError:
                    salaire_flt = 0.0

                cursor.execute('''
                    INSERT INTO employes (matricule, nom, prenom, salaire_base, statut, id_site_affecte, equipe)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                ''', (matricule, nom, prenom, salaire_flt, statut.capitalize(), id_site, equipe))
                nb_importes += 1

        else:
            equipe_courante = "MATIN"
            
            for ligne in lignes_donnees:
                texte_ligne = " ".join([str(c) for c in ligne if c]).strip()
                
                if "equipe soir" in texte_ligne.lower() or "équipe soir" in texte_ligne.lower():
                    equipe_courante = "SOIR"
                elif "equipe matin" in texte_ligne.lower() or "équipe matin" in texte_ligne.lower():
                    equipe_courante = "MATIN"

                match = re.search(r'^\s*(\d+)?\s*([A-Za-zÀ-ÿ\'\-]+)\s+([A-Za-zÀ-ÿ\'\-\s]+)$', texte_ligne)
                
                if match:
                    p1 = match.group(2).strip()
                    p2 = match.group(3).strip()
                    
                    mots_ignores = ['vendredi', 'samedi', 'dimanche', 'lundi', 'mardi', 'mercredi', 'jeudi', 'entrée', 'sortie', 'fiche', 'pointage', 'note', 'admin']
                    if any(m.lower() in p1.lower() or m.lower() in p2.lower() for m in mots_ignores):
                        continue

                    prenom = p1
                    nom = p2
                    
                    cursor.execute("SELECT COUNT(*) FROM employes")
                    total_existants = cursor.fetchone()[0]
                    matricule = f"EMP-{total_existants + nb_importes + 1:03d}"

                    cursor.execute("SELECT matricule FROM employes WHERE LOWER(nom) = %s AND LOWER(prenom) = %s", (nom.lower(), prenom.lower()))
                    if cursor.fetchone():
                        continue

                    id_site = get_or_create_site(cursor, "BFI")

                    cursor.execute('''
                        INSERT INTO employes (matricule, nom, prenom, salaire_base, statut, id_site_affecte, equipe)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                    ''', (matricule, nom, prenom, 15000.0, 'Actif', id_site, equipe_courante))
                    nb_importes += 1

        conn.commit()
        cursor.close()
        conn.close()

        return f"<h3>Succès : {nb_importes} employé(s) enregistré(s) !</h3><br><a href='/'>Retour au tableau de bord</a>"

    except Exception as e:
        print(f"Erreur d'importation : {e}")
        return f"<h3>Erreur lors de la lecture du fichier :</h3><p>{e}</p><br><a href='/'>Retour</a>"


def get_or_create_site(cursor, nom_site):
    if not nom_site:
        return None
    cursor.execute("SELECT id FROM sites WHERE LOWER(nom) = %s", (nom_site.lower(),))
    res = cursor.fetchone()
    if res:
        return res[0]
    else:
        cursor.execute("INSERT INTO sites (nom, adresse) VALUES (%s, %s)", (nom_site, "Créé via import auto"))
        try:
            cursor.execute("SELECT LASTVAL()")
        except Exception:
            cursor.execute("SELECT LAST_INSERT_ID()")
        return cursor.fetchone()[0]


@app.route('/importer_paie', methods=['POST'])
def importer_paie():
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    file = request.files.get('file_paie')
    if not file or file.filename == '':
        return f"<h3>Aucun fichier sélectionné.</h3><br><a href='/paie'>Retour</a>"

    filename = file.filename.lower()
    lignes_donnees = []

    try:
        if filename.endswith(('.xlsx', '.xls', '.csv')):
            try:
                if filename.endswith('.csv'):
                    df = pd.read_csv(file, sep=None, engine='python', dtype=str)
                else:
                    df = pd.read_excel(file, dtype=str)
                df = df.fillna('')
                lignes_donnees = [df.columns.astype(str).tolist()] + df.values.tolist()
            except Exception:
                file.seek(0)
                content = file.stream.read().decode("utf-8-sig", errors='ignore')
                stream = io.StringIO(content)
                separateur = ';' if ';' in content.split('\n')[0] else ','
                reader = csv.reader(stream, delimiter=separateur)
                lignes_donnees = list(reader)

        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        lignes_donnees.append(cells)

        elif filename.endswith('.pdf'):
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if any(row):
                                lignes_donnees.append([str(cell or '').strip() for cell in row])

        if not lignes_donnees:
            return f"<h3>Aucune donnée extraite du fichier.</h3><br><a href='/paie'>Retour</a>"

        headers = [str(h).lower().strip() for h in lignes_donnees[0]]

        idx_mat = next((i for i, h in enumerate(headers) if 'matricule' in h), 0)
        idx_salaire = next((i for i, h in enumerate(headers) if 'salaire' in h or 'base' in h), -1)

        conn = get_db_connection()
        cursor = conn.cursor()
        nb_majd = 0

        for row in lignes_donnees[1:]:
            if len(row) <= idx_mat:
                continue

            matricule = str(row[idx_mat]).upper().strip()
            if not matricule:
                continue

            if idx_salaire != -1 and idx_salaire < len(row):
                try:
                    sal_str = str(row[idx_salaire]).replace(',', '.').replace(' ', '')
                    nouveau_salaire = float(sal_str)
                    cursor.execute("UPDATE employes SET salaire_base = %s WHERE matricule = %s", (nouveau_salaire, matricule))
                    nb_majd += 1
                except ValueError:
                    pass

        conn.commit()
        cursor.close()
        conn.close()

        return f"<h3>Importation de la paie réussie ! {nb_majd} employé(s) mis à jour.</h3><br><a href='/paie'>Retour au Rapport de Paie</a>"

    except Exception as e:
        print(f"Erreur d'importation Paie : {e}")
        return f"<h3>Erreur lors de l'importation de la paie :</h3><p>{e}</p><br><a href='/paie'>Retour</a>"


@app.route('/supprimer_employe/<matricule>', methods=['POST'])
def supprimer_employe(matricule):
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM pointages WHERE matricule_employe = %s", (matricule,))
    cursor.execute("DELETE FROM employes WHERE matricule = %s", (matricule,))
    conn.commit()
    cursor.close()
    conn.close()
    return redirect(url_for('dashboard'))


@app.route('/modifier_equipe_employe/<matricule>', methods=['POST'])
def modifier_equipe_employe(matricule):
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    nouvelle_equipe = request.form.get('equipe', 'MATIN').upper()

    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        cursor.execute('''
            UPDATE employes 
            SET equipe = %s 
            WHERE matricule = %s
        ''', (nouvelle_equipe, matricule))
        
        conn.commit()
    except Exception as e:
        print(f"Erreur lors du changement d'équipe : {e}")
        conn.rollback()
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('dashboard'))


@app.route('/paie')
def rapport_paie():
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT matricule, nom, prenom, salaire_base FROM employes")
    liste_employes = cursor.fetchall()
    
    bilan_paie = []
    for emp in liste_employes:
        matricule, nom, prenom, salaire_base = emp
        cursor.execute('SELECT heure_arrivee, heure_depart FROM pointages WHERE matricule_employe = %s AND heure_depart IS NOT NULL', (matricule,))
        pointages = cursor.fetchall()
        
        total_heures = 0.0
        for p in pointages:
            total_heures += calculer_heures(p[0], p[1])
            
        taux_horaire = round(salaire_base / 160, 2)
        salaire_gagne = round(total_heures * taux_horaire, 2)
        
        bilan_paie.append({
            'matricule': matricule, 'nom': nom, 'prenom': prenom,
            'heures': total_heures, 'taux': taux_horaire, 'salaire_du': salaire_gagne
        })
        
    cursor.close()
    conn.close()
    return render_template('paie.html', bilan=bilan_paie)


@app.route('/exporter_paie_csv')
def exporter_paie_csv():
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT matricule, nom, prenom, salaire_base FROM employes")
    liste_employes = cursor.fetchall()
    
    output = io.StringIO()
    writer = csv.writer(output, delimiter=';') 
    
    writer.writerow(['Matricule', 'Employe', 'Total Heures', 'Taux Horaire (MRU/h)', 'Salaire a Verser (MRU)'])
    
    for emp in liste_employes:
        matricule, nom, prenom, salaire_base = emp
        cursor.execute('SELECT heure_arrivee, heure_depart FROM pointages WHERE matricule_employe = %s AND heure_depart IS NOT NULL', (matricule,))
        pointages = cursor.fetchall()
        
        total_heures = 0.0
        for p in pointages:
            total_heures += calculer_heures(p[0], p[1])
            
        taux_horaire = round(salaire_base / 160, 2)
        salaire_gagne = round(total_heures * taux_horaire, 2)
        
        nom_complet = f"{prenom} {nom}"
        writer.writerow([matricule, nom_complet, f"{total_heures} h", f"{taux_horaire} MRU", f"{salaire_gagne} MRU"])
        
    cursor.close()
    conn.close()
    
    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-disposition": "attachment; filename=Rapport_Paie_Nettoyage.csv"}
    )


@app.route('/supprimer_site/<int:id>', methods=['POST'])
def supprimer_site(id):
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM pointages WHERE id_site = %s", (id,))
        cursor.execute("DELETE FROM sites WHERE id = %s", (id,))
        conn.commit()
    except Exception as e:
        print(f"Erreur lors de la suppression : {e}")
        conn.rollback()
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for('dashboard'))


@app.route('/modifier_site/<int:id>', methods=['GET', 'POST'])
def modifier_site(id):
    if not session.get('est_admin'):
        return redirect(url_for('espace_pointage'))

    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        nouveau_nom = request.form.get('nom_site')
        nouvelle_adresse = request.form.get('adresse') or request.form.get('adresse_site')
        
        cursor.execute("UPDATE sites SET nom = %s, adresse = %s WHERE id = %s", 
                       (nouveau_nom, nouvelle_adresse, id))
        conn.commit()
        cursor.close()
        conn.close()
        return redirect(url_for('dashboard'))
    
    cursor.execute("SELECT id, nom, adresse FROM sites WHERE id = %s", (id,))
    site_data = cursor.fetchone()
    cursor.close()
    conn.close()
    
    if not site_data:
        return redirect(url_for('dashboard'))
        
    site = {'id': site_data[0], 'nom_site': site_data[1], 'adresse': site_data[2]}
    return render_template('modifier_site.html', site=site)


# --- ESPACE CHEF D'ÉQUIPE (LIBRE D'ACCÈS) ---

@app.route('/pointage')
def espace_pointage():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    site_id_selected = request.args.get('site_id')
    
    cursor.execute("SELECT id, nom FROM sites")
    liste_sites = cursor.fetchall()
    
    if not site_id_selected and liste_sites:
        site_id_selected = liste_sites[0][0]
        
    if site_id_selected:
        cursor.execute("""
            SELECT matricule, nom, prenom, COALESCE(equipe, 'MATIN') as equipe 
            FROM employes 
            WHERE statut = 'Actif' AND (id_site_affecte = %s OR id_site_affecte IS NULL)
        """, (int(site_id_selected),))
    else:
        cursor.execute("""
            SELECT matricule, nom, prenom, COALESCE(equipe, 'MATIN') as equipe 
            FROM employes 
            WHERE statut = 'Actif'
        """)
    
    tous_employes = cursor.fetchall()
    
    employes_matin = [e for e in tous_employes if e[3].upper() == 'MATIN']
    employes_soir = [e for e in tous_employes if e[3].upper() == 'SOIR']
    
    cursor.close()
    conn.close()
    
    return render_template(
        'pointage.html', 
        sites=liste_sites, 
        site_selected=int(site_id_selected) if site_id_selected else None,
        employes_matin=employes_matin,
        employes_soir=employes_soir
    )

@app.route('/executer_pointage', methods=['POST'])
def executer_pointage():
    matricule = request.form.get('matricule')
    site_id = request.form.get('site_id')
    equipe_form = request.form.get('equipe')
    action = request.form.get('action')
    lat = request.form.get('latitude')
    lng = request.form.get('longitude')
    fiche_id = request.form.get('fiche_id')
    agent = request.form.get('agent')
    observations = request.form.get('observations')

    if not lat or lat.strip() == "": lat = None
    if not lng or lng.strip() == "": lng = None

    date_aujourdhui = datetime.now().strftime('%Y-%m-%d')
    heure_actuelle = datetime.now().strftime('%H:%M:%S')
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        if not equipe_form:
            cursor.execute("SELECT COALESCE(equipe, 'MATIN') FROM employes WHERE matricule = %s", (matricule,))
            res = cursor.fetchone()
            equipe = res[0] if res else 'MATIN'
        else:
            equipe = equipe_form.upper().strip()

        if action == 'arrivee':
            cursor.execute('''
                INSERT INTO pointages (matricule_employe, id_site, date_jour, heure_arrivee, latitude, longitude, equipe, fiche_id, agent, observations)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''', (matricule, int(site_id), date_aujourdhui, heure_actuelle, lat, lng, equipe, fiche_id, agent, observations))

        elif action == 'depart':
            cursor.execute('''
                UPDATE pointages 
                SET heure_depart = %s 
                WHERE matricule_employe = %s AND date_jour = %s AND heure_depart IS NULL
            ''', (heure_actuelle, matricule, date_aujourdhui))
            
            if cursor.rowcount == 0:
                if IS_RENDER:
                    cursor.execute('''
                        UPDATE pointages 
                        SET heure_depart = %s 
                        WHERE id = (
                            SELECT id FROM pointages 
                            WHERE matricule_employe = %s AND heure_depart IS NULL 
                            ORDER BY id DESC LIMIT 1
                        )
                    ''', (heure_actuelle, matricule))
                else:
                    cursor.execute('''
                        UPDATE pointages 
                        SET heure_depart = %s 
                        WHERE id = (
                            SELECT id FROM (
                                SELECT id FROM pointages 
                                WHERE matricule_employe = %s AND heure_depart IS NULL 
                                ORDER BY id DESC LIMIT 1
                            ) as t
                        )
                    ''', (heure_actuelle, matricule))
                
        conn.commit()
        message = "<h3>Pointage réussi ! Merci.</h3>"
    except Exception as e:
        conn.rollback()
        message = f"<h3>Une erreur est survenue lors de l'enregistrement.</h3><p>{e}</p>"
    finally:
        cursor.close()
        conn.close()
        
    return f"{message}<br><a href='/pointage?site_id={site_id}'>Retour</a>"
    
def calculer_heures(arrivee, depart):
    if not arrivee or not depart:
        return 0.0
    fmt = '%H:%M:%S'
    t_arrivee = datetime.strptime(arrivee, fmt)
    t_depart = datetime.strptime(depart, fmt)
    
    diff = t_depart - t_arrivee
    if diff.total_seconds() < 0:
        diff_seconds = diff.total_seconds() + 86400
    else:
        diff_seconds = diff.total_seconds()
        
    return round(diff_seconds / 3600, 2)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)